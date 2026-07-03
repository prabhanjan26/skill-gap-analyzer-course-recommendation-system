"""Stage 1: Resume parsing (DDD 6 Stage 1).

Extracts raw text from PDF (PyPDF2) or DOCX (python-docx), cleans whitespace,
strips obvious page numbers, and enforces a 100-char minimum.
"""
from __future__ import annotations

import logging
import os
import re

logger = logging.getLogger(__name__)

MIN_TEXT_LENGTH = 100

PDF_MIMES = {"application/pdf"}
DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx", ".doc"}


class ResumeParseError(Exception):
    """Base class for resume parsing errors."""


class UnsupportedFileTypeError(ResumeParseError):
    pass


class CorruptedFileError(ResumeParseError):
    pass


class InsufficientContentError(ResumeParseError):
    pass


def _clean_text(text: str) -> str:
    """Collapse whitespace and remove standalone page-number lines."""
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        # Drop lines that are just a page number, e.g. "3" or "Page 3 of 8".
        if re.fullmatch(r"\d{1,3}", line):
            continue
        if re.fullmatch(r"(?i)page\s+\d+\s+of\s+\d+", line):
            continue
        # Collapse internal runs of whitespace.
        line = re.sub(r"[ \t]+", " ", line)
        lines.append(line)
    return "\n".join(lines).strip()


def _extract_pdf(file_path: str) -> tuple[str, int]:
    from PyPDF2 import PdfReader

    try:
        reader = PdfReader(file_path)
        pages = reader.pages
        texts = []
        for page in pages:
            try:
                texts.append(page.extract_text() or "")
            except Exception:  # pragma: no cover - per-page resilience
                continue
        return "\n".join(texts), len(pages)
    except Exception as exc:  # pragma: no cover
        raise CorruptedFileError(f"Failed to read PDF: {exc}") from exc


def _extract_docx(file_path: str) -> tuple[str, int]:
    import docx

    try:
        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs]
        # Include table cell text (skills are often tabular).
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs), 1
    except Exception as exc:  # pragma: no cover
        raise CorruptedFileError(f"Failed to read DOCX: {exc}") from exc


def _detect_kind(file_path: str, mime_type: str | None) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    mime = (mime_type or "").lower()
    if mime in PDF_MIMES or ext in PDF_EXTS:
        return "pdf"
    if mime in DOCX_MIMES or ext in DOCX_EXTS:
        return "docx"
    raise UnsupportedFileTypeError("Invalid file type. Accepted: PDF, DOCX")


def parse_resume(file_path: str, mime_type: str | None = None) -> dict:
    """Parse a resume file into cleaned text.

    Returns { raw_text, page_count, metadata }.
    Raises UnsupportedFileTypeError / CorruptedFileError / InsufficientContentError.
    """
    if not os.path.exists(file_path):
        raise CorruptedFileError("Uploaded file could not be found on disk")

    if os.path.getsize(file_path) == 0:
        raise InsufficientContentError("Empty file uploaded")

    kind = _detect_kind(file_path, mime_type)
    raw, page_count = _extract_pdf(file_path) if kind == "pdf" else _extract_docx(file_path)

    cleaned = _clean_text(raw)
    if len(cleaned) < MIN_TEXT_LENGTH:
        raise InsufficientContentError(
            "Insufficient text content (resume produced fewer than "
            f"{MIN_TEXT_LENGTH} characters after cleaning)"
        )

    return {
        "raw_text": cleaned,
        "page_count": page_count,
        "metadata": {"kind": kind, "char_count": len(cleaned)},
    }
